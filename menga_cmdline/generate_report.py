from docx import Document
from docx.shared import Inches
import datetime
import cairosvg

#document.add_picture('monty-truth.png', width=Inches(1.25))

def add_setup_information(document, setup):
    document.add_heading('Setup Information', level=1)
    
    document.add_paragraph('')
    document.add_paragraph('The information of the environment used to make the analysis are the following')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    timestamp='{:%Y-%m-%d %H:%M:%S}'.format(datetime.datetime.now())
    records = (
    ('Image Name', setup[0]),
    ('Main PID', setup[1]),
    ('Analysis Time', setup[2]+'s'),
    ('Timestamp', timestamp)
    )

    table = document.add_table(rows=0, cols=2)
    for desc, value in records:
        row_cells = table.add_row().cells
        row_cells[0].text = desc
        row_cells[1].text = value

    document.add_page_break()

    return document

def add_network_analysis(document, path):
    document.add_heading('Network Activity', level=1)
    document.add_paragraph('')
    document.add_paragraph('')
    table = document.add_table(rows=0, cols=7)

    with open(path,'r') as f:
        for line in f:
            data=line.split(';')
            row_cells = table.add_row().cells

            timestamp =data[0]
            try:
                timestamp='{:%Y-%m-%d %H:%M:%S}'.format(datetime.datetime.fromtimestamp(float(data[0])))
            except ValueError:
                timestamp = data[0]
            
            row_cells[0].text = timestamp
            row_cells[1].text = data[1]
            row_cells[2].text = data[2]
            row_cells[3].text = data[3]
            row_cells[4].text = data[4]
            row_cells[5].text = data[5]
            row_cells[6].text = data[6]

    document.add_page_break()

    return document

def add_cpu_analysis(document, path):
    document.add_heading('CPU Activity', level=1)
    document.add_paragraph('')
    document.add_paragraph('The graph represent the cpu consumption with its name on each bar')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    cairosvg.svg2png(url=path,write_to=path+'.png')
    document.add_picture(path+'.png', width=Inches(7))
    document.add_page_break()
    return document

def add_kernel_analysis(document, path):
    document.add_heading('Kernel Activity', level=1)
    document.add_paragraph('')
    document.add_paragraph('')
    #print(path)
    table = document.add_table(rows=0, cols=6)
    with open(path,'r') as f:
        for line in f:
            data=line.split(';')
            #print(str(data))
            row_cells = table.add_row().cells

            timestamp =data[0]
            try:
                timestamp='{:%Y-%m-%d %H:%M:%S}'.format(datetime.datetime.fromtimestamp(float(data[0])))
            except ValueError:
                timestamp = data[0]
            
            row_cells[0].text = timestamp
            row_cells[1].text = data[1]
            row_cells[2].text = data[2]
            row_cells[3].text = data[3]
            row_cells[4].text = data[4]
    
    document.add_page_break()
    return document

def generate_menga_report(output, setup, network_path, cpu_path, kernel_path):
    document = Document("./menga_cmdline/template-menga.docx")
    document = add_setup_information(document,setup)
    add_network_analysis(document, network_path)
    add_cpu_analysis(document, cpu_path)
    add_kernel_analysis(document, kernel_path)
    document.save(output)


